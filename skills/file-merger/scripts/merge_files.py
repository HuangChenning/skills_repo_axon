#!/usr/bin/env python3
"""
文件批量合并脚本（增强版）
支持 Markdown、PDF、Word、TXT、EPUB、HTML、RTF、ODT 文件的批量合并
扩展功能：递归扫描、正则过滤、大小平衡、报告生成、增量合并
"""

import os
import sys
import logging
import argparse
import re
import hashlib
import csv
from pathlib import Path
from typing import List, Dict, Optional, Tuple, Set
from datetime import datetime
from dataclasses import dataclass, asdict
import json

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def natural_sort_key(text):
    """
    自然排序键函数

    将字符串中的数字部分转换为整数，实现正确的数字排序
    例如: chapter_1.md < chapter_10.md < chapter_2.md

    Args:
        text: 字符串（通常是文件名）

    Returns:
        用于排序的元组
    """
    def tryint(s):
        try:
            return int(s)
        except ValueError:
            return s

    return [tryint(c) for c in re.split(r'(\d+)', str(text))]


@dataclass
class SourceFile:
    """源文件信息"""
    path: Path
    name: str
    size: int
    hash: str


@dataclass
class OutputFile:
    """输出文件信息"""
    name: str
    source_files: List[SourceFile]
    total_size: int


@dataclass
class MergeReport:
    """合并报告"""
    timestamp: str
    input_directory: str
    output_directory: str
    file_type: str
    scan_mode: str
    total_files_found: int
    total_files_processed: int
    outputs_created: int
    balance_mode: str
    outputs: List[OutputFile]
    errors: List[str]


class FileMerger:
    """文件合并器（增强版）"""

    # 支持的文件类型
    FILE_TYPES = {
        'md': {'ext': '.md', 'name': 'Markdown', 'mime': 'text/markdown'},
        'pdf': {'ext': '.pdf', 'name': 'PDF', 'mime': 'application/pdf'},
        'docx': {'ext': '.docx', 'name': 'Word', 'mime': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'},
        'txt': {'ext': '.txt', 'name': '文本', 'mime': 'text/plain'},
        'epub': {'ext': '.epub', 'name': 'EPUB', 'mime': 'application/epub+zip'},
        'html': {'ext': '.html', 'name': 'HTML', 'mime': 'text/html'},
        'htm': {'ext': '.htm', 'name': 'HTML', 'mime': 'text/html'},
        'rtf': {'ext': '.rtf', 'name': 'RTF', 'mime': 'application/rtf'},
        'odt': {'ext': '.odt', 'name': 'ODT', 'mime': 'application/vnd.oasis.opendocument.text'}
    }

    def __init__(
        self,
        input_dir: str,
        output_dir: str,
        file_type: str,
        output_count: int,
        order_file: Optional[str] = None,
        total_files: Optional[int] = None,
        separator: str = '---',
        generate_index: bool = True,
        recursive: bool = False,
        filter_pattern: Optional[str] = None,
        exclude_pattern: Optional[str] = None,
        balance_by_size: bool = False,
        generate_report: Optional[str] = None,
        incremental: bool = False,
        state_file: Optional[str] = None
    ):
        """
        初始化文件合并器（增强版）

        Args:
            input_dir: 输入目录路径
            output_dir: 输出目录路径
            file_type: 文件类型 (md/pdf/docx/txt/epub/html/rtf/odt/all)
            output_count: 要生成的输出文件数量
            order_file: 自定义排序文件路径
            total_files: 指定要处理的文件总数
            separator: 文件分隔符
            generate_index: 是否生成目录索引
            recursive: 是否递归扫描子目录
            filter_pattern: 包含文件的正则表达式
            exclude_pattern: 排除文件的正则表达式
            balance_by_size: 按文件大小平衡分配
            generate_report: 生成报告文件路径（JSON/CSV）
            incremental: 增量合并模式
            state_file: 状态文件路径（用于增量合并）
        """
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        self.file_type = file_type
        self.output_count = output_count
        self.order_file = order_file
        self.total_files = total_files
        self.separator = separator
        self.generate_index = generate_index
        self.recursive = recursive
        self.filter_pattern = re.compile(filter_pattern) if filter_pattern else None
        self.exclude_pattern = re.compile(exclude_pattern) if exclude_pattern else None
        self.balance_by_size = balance_by_size
        self.generate_report = generate_report
        self.incremental = incremental
        self.state_file = state_file

        # 验证输入目录
        if not self.input_dir.exists():
            raise ValueError(f"输入目录不存在: {input_dir}")

        # 创建输出目录
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # 统计信息
        self.stats = {
            'total_scanned': 0,
            'total_processed': 0,
            'outputs_created': 0,
            'errors': []
        }

        # 报告数据
        self.report_outputs: List[OutputFile] = []

        # 状态文件数据
        self.processed_files: Set[str] = set()
        if self.incremental and self.state_file:
            self._load_state()

    def _load_state(self):
        """加载状态文件（用于增量合并）"""
        state_path = Path(self.state_file)
        if state_path.exists():
            try:
                with open(state_path, 'r', encoding='utf-8') as f:
                    state = json.load(f)
                    self.processed_files = set(state.get('processed_files', []))
                logger.info(f"加载状态文件: {len(self.processed_files)} 个已处理文件")
            except Exception as e:
                logger.error(f"加载状态文件失败: {e}")

    def _save_state(self):
        """保存状态文件"""
        if not self.state_file:
            return

        try:
            state = {
                'timestamp': datetime.now().isoformat(),
                'processed_files': list(self.processed_files),
                'last_update': datetime.now().isoformat()
            }
            with open(self.state_file, 'w', encoding='utf-8') as f:
                json.dump(state, f, ensure_ascii=False, indent=2)
            logger.info(f"状态已保存: {self.state_file}")
        except Exception as e:
            logger.error(f"保存状态文件失败: {e}")

    def _calculate_file_hash(self, file_path: Path) -> str:
        """计算文件的哈希值"""
        try:
            hasher = hashlib.md5()
            with open(file_path, 'rb') as f:
                # 只读取前 8KB 来计算哈希（提高速度）
                chunk = f.read(8192)
                hasher.update(chunk)
            return hasher.hexdigest()
        except Exception as e:
            logger.warning(f"计算文件哈希失败 {file_path}: {e}")
            return ""

    def scan_files(self) -> Dict[str, List[SourceFile]]:
        """
        扫描输入目录中的文件（支持递归和过滤）

        Returns:
            按文件类型分组的文件信息字典
        """
        logger.info(f"正在扫描目录: {self.input_dir} (递归: {self.recursive})")

        files_by_type = {}

        # 确定要处理的文件类型
        types_to_process = []
        if self.file_type == 'all':
            types_to_process = list(self.FILE_TYPES.keys())
        elif self.file_type in self.FILE_TYPES:
            types_to_process = [self.file_type]
        else:
            raise ValueError(f"不支持的文件类型: {self.file_type}")

        # 扫描文件
        glob_pattern = "**/*" if self.recursive else "*"

        for file_type in types_to_process:
            ext = self.FILE_TYPES[file_type]['ext']
            files = []

            # 扩展 HTML 支持两个扩展名
            if file_type == 'html':
                paths = list(self.input_dir.glob(f"**/*.html")) + list(self.input_dir.glob(f"**/*.htm"))
            else:
                paths = list(self.input_dir.glob(f"{glob_pattern}{ext}"))

            for file_path in paths:
                # 跳过目录
                if file_path.is_dir():
                    continue

                # 应用包含过滤
                if self.filter_pattern and not self.filter_pattern.search(file_path.name):
                    continue

                # 应用排除过滤
                if self.exclude_pattern and self.exclude_pattern.search(file_path.name):
                    logger.debug(f"排除文件: {file_path.name}")
                    continue

                # 增量模式：跳过已处理的文件
                if self.incremental:
                    file_hash = self._calculate_file_hash(file_path)
                    file_key = f"{file_path}:{file_hash}"
                    if file_key in self.processed_files:
                        logger.debug(f"跳过已处理文件: {file_path}")
                        continue

                # 获取文件信息
                try:
                    stat = file_path.stat()
                    source_file = SourceFile(
                        path=file_path,
                        name=file_path.name,
                        size=stat.st_size,
                        hash=self._calculate_file_hash(file_path)
                    )
                    files.append(source_file)
                except Exception as e:
                    logger.warning(f"无法读取文件信息 {file_path}: {e}")

            if files:
                # 使用自然排序
                files.sort(key=lambda f: natural_sort_key(f.name))
                files_by_type[file_type] = files
                logger.info(f"找到 {len(files)} 个 {ext} 文件")
            else:
                logger.warning(f"未找到 {ext} 文件")

        self.stats['total_scanned'] = sum(len(files) for files in files_by_type.values())

        return files_by_type

    def load_order_file(self) -> Optional[List[str]]:
        """
        加载排序文件

        Returns:
            排序后的文件名列表
        """
        if not self.order_file:
            return None

        order_path = Path(self.order_file)
        if not order_path.exists():
            logger.error(f"排序文件不存在: {self.order_file}")
            return None

        logger.info(f"正在加载排序文件: {self.order_file}")

        order = []
        with open(order_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                # 跳过空行和注释
                if not line or line.startswith('#'):
                    continue
                order.append(line)

        logger.info(f"加载了 {len(order)} 个排序项")
        return order

    def sort_files(self, files: List[SourceFile], order: Optional[List[str]] = None) -> List[SourceFile]:
        """
        对文件列表进行排序

        Args:
            files: 文件信息列表
            order: 自定义排序列表

        Returns:
            排序后的文件列表
        """
        if order:
            # 按自定义顺序排序
            logger.info("使用自定义排序")
            file_dict = {f.name: f for f in files}
            sorted_files = []

            for name in order:
                if name in file_dict:
                    sorted_files.append(file_dict[name])
                else:
                    logger.warning(f"排序文件中指定的文件不存在: {name}")

            # 添加未在排序文件中的文件
            remaining = [f for f in files if f.name not in order]
            remaining.sort(key=lambda f: natural_sort_key(f.name))
            sorted_files.extend(remaining)

            return sorted_files
        else:
            # 已在扫描时排序，直接返回
            return files

    def calculate_allocation(self, files: List[SourceFile]) -> List[tuple]:
        """
        计算文件分配方案

        Args:
            files: 文件信息列表

        Returns:
            [(start_idx, end_idx), ...] 分配列表
        """
        if not files:
            return []

        if self.balance_by_size:
            return self._calculate_allocation_by_size(files)
        else:
            return self._calculate_allocation_by_count(files)

    def _calculate_allocation_by_count(self, files: List[SourceFile]) -> List[tuple]:
        """
        按文件数量均匀分配

        Args:
            files: 文件信息列表

        Returns:
            分配列表
        """
        file_count = len(files)

        # 计算每个输出文件应包含的源文件数
        files_per_output = file_count // self.output_count
        remainder = file_count % self.output_count

        allocation = []
        start = 0

        for i in range(self.output_count):
            # 前面几个输出多分配一个文件（如果有余数）
            extra = 1 if i < remainder else 0
            end = start + files_per_output + extra

            # 确保不超过文件总数
            if start >= file_count:
                break

            end = min(end, file_count)
            allocation.append((start, end))

            start = end

        return allocation

    def _calculate_allocation_by_size(self, files: List[SourceFile]) -> List[tuple]:
        """
        按文件大小平衡分配

        Args:
            files: 文件信息列表

        Returns:
            分配列表
        """
        total_size = sum(f.size for f in files)
        target_size = total_size // self.output_count

        logger.info(f"总大小: {total_size} 字节, 每个输出目标大小: {target_size} 字节")

        allocation = []
        current_size = 0
        start_idx = 0

        for i in range(self.output_count):
            if start_idx >= len(files):
                break

            # 从当前索引开始累积文件
            current_size = 0
            end_idx = start_idx

            while end_idx < len(files):
                # 如果添加下一个文件会超过目标大小太多，则停止
                if current_size + files[end_idx].size > target_size * 1.5 and current_size > 0:
                    break
                current_size += files[end_idx].size
                end_idx += 1

            # 确保至少有一个文件
            if end_idx == start_idx and end_idx < len(files):
                end_idx = start_idx + 1

            allocation.append((start_idx, end_idx))
            logger.info(f"输出 {i + 1}: {end_idx - start_idx} 个文件, 总大小: {current_size} 字节")

            start_idx = end_idx

        return allocation

    def merge_markdown(self, files: List[SourceFile], output_path: Path) -> bool:
        """
        合并 Markdown 文件

        Args:
            files: 源文件信息列表
            output_path: 输出文件路径

        Returns:
            成功返回 True
        """
        try:
            content_parts = []

            # 生成目录索引
            if self.generate_index:
                index = "# 目录\n\n"
                for i, source_file in enumerate(files, 1):
                    index += f"{i}. {source_file.path.stem}\n"
                content_parts.append(index)
                content_parts.append(f"\n{self.separator}\n\n")

            # 合并文件内容
            for source_file in files:
                try:
                    with open(source_file.path, 'r', encoding='utf-8') as f:
                        content = f.read()

                    # 添加文件标题
                    content_parts.append(f"## {source_file.path.stem}\n\n")
                    content_parts.append(content)
                    content_parts.append(f"\n\n{self.separator}\n\n")

                except Exception as e:
                    logger.error(f"读取文件失败 {source_file.path}: {e}")
                    self.stats['errors'].append(str(source_file.path))

            # 写入输出文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.writelines(content_parts)

            logger.info(f"已保存: {output_path}")
            return True

        except Exception as e:
            logger.error(f"合并 Markdown 失败: {e}")
            return False

    def merge_text(self, files: List[SourceFile], output_path: Path) -> bool:
        """
        合并纯文本文件

        Args:
            files: 源文件信息列表
            output_path: 输出文件路径

        Returns:
            成功返回 True
        """
        try:
            import chardet

            content_parts = []

            # 生成目录索引
            if self.generate_index:
                index = "# 目录\n\n"
                for i, source_file in enumerate(files, 1):
                    index += f"{i}. {source_file.name}\n"
                content_parts.append(index)
                content_parts.append(f"\n{self.separator}\n\n")

            # 合并文件内容
            for source_file in files:
                try:
                    # 检测编码
                    with open(source_file.path, 'rb') as f:
                        raw_data = f.read()
                        encoding_result = chardet.detect(raw_data)
                        encoding = encoding_result['encoding'] or 'utf-8'

                    # 读取内容
                    with open(source_file.path, 'r', encoding=encoding) as f:
                        content = f.read()

                    content_parts.append(f"## {source_file.name}\n\n")
                    content_parts.append(content)
                    content_parts.append(f"\n\n{self.separator}\n\n")

                except Exception as e:
                    logger.error(f"读取文件失败 {source_file.path}: {e}")
                    self.stats['errors'].append(str(source_file.path))

            # 写入输出文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.writelines(content_parts)

            logger.info(f"已保存: {output_path}")
            return True

        except ImportError:
            logger.error("未安装 chardet 库，请运行: pip3 install chardet")
            return False
        except Exception as e:
            logger.error(f"合并文本文件失败: {e}")
            return False

    def merge_pdf(self, files: List[SourceFile], output_path: Path) -> bool:
        """
        合并 PDF 文件

        Args:
            files: 源文件信息列表
            output_path: 输出文件路径

        Returns:
            成功返回 True
        """
        try:
            from PyPDF2 import PdfMerger

            merger = PdfMerger()

            for source_file in files:
                try:
                    merger.append(str(source_file.path))
                    logger.info(f"已添加: {source_file.name}")
                except Exception as e:
                    logger.error(f"添加 PDF 失败 {source_file.path}: {e}")
                    self.stats['errors'].append(str(source_file.path))

            # 写入输出文件
            merger.write(str(output_path))
            merger.close()

            logger.info(f"已保存: {output_path}")
            return True

        except ImportError:
            logger.error("未安装 PyPDF2 库，请运行: pip3 install pypdf2")
            return False
        except Exception as e:
            logger.error(f"合并 PDF 失败: {e}")
            return False

    def merge_docx(self, files: List[SourceFile], output_path: Path) -> bool:
        """
        合并 Word 文档

        Args:
            files: 源文件信息列表
            output_path: 输出文件路径

        Returns:
            成功返回 True
        """
        try:
            from docx import Document

            # 创建新文档
            merged_doc = Document()

            # 添加目录索引
            if self.generate_index:
                merged_doc.add_heading('目录', 0)
                for i, source_file in enumerate(files, 1):
                    merged_doc.add_paragraph(f"{i}. {source_file.path.stem}")
                merged_doc.add_paragraph(self.separator)

            # 合并文档内容
            for source_file in files:
                try:
                    doc = Document(str(source_file.path))

                    # 添加文件标题
                    merged_doc.add_heading(source_file.path.stem, 1)

                    # 复制所有段落
                    for paragraph in doc.paragraphs:
                        merged_doc.add_paragraph(paragraph.text)

                    # 复制所有表格
                    for table in doc.tables:
                        merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                        new_table = merged_doc.tables[-1]
                        for i, row in enumerate(table.rows):
                            for j, cell in enumerate(row.cells):
                                new_table.rows[i].cells[j].text = cell.text

                    # 添加分隔符
                    merged_doc.add_paragraph(self.separator)

                except Exception as e:
                    logger.error(f"读取 Word 文档失败 {source_file.path}: {e}")
                    self.stats['errors'].append(str(source_file.path))

            # 保存合并后的文档
            merged_doc.save(str(output_path))

            logger.info(f"已保存: {output_path}")
            return True

        except ImportError:
            logger.error("未安装 python-docx 库，请运行: pip3 install python-docx")
            return False
        except Exception as e:
            logger.error(f"合并 Word 文档失败: {e}")
            return False

    def merge_html(self, files: List[SourceFile], output_path: Path) -> bool:
        """
        合并 HTML 文件

        Args:
            files: 源文件信息列表
            output_path: 输出文件路径

        Returns:
            成功返回 True
        """
        try:
            from bs4 import BeautifulSoup

            # 创建主文档
            soup = BeautifulSoup('<html><head><meta charset="UTF-8"><title>Merged Document</title></head><body></body></html>', 'html.parser')

            body = soup.find('body')

            # 添加目录索引
            if self.generate_index:
                index_div = soup.new_tag('div', **{'class': 'index'})
                index_h1 = soup.new_tag('h1')
                index_h1.string = '目录'
                index_div.append(index_h1)

                index_ul = soup.new_tag('ul')
                for i, source_file in enumerate(files, 1):
                    index_li = soup.new_tag('li')
                    index_a = soup.new_tag('a', href=f"#section-{i}")
                    index_a.string = source_file.path.stem
                    index_li.append(index_a)
                    index_ul.append(index_li)

                index_div.append(index_ul)
                body.append(index_div)

                hr = soup.new_tag('hr')
                body.append(hr)

            # 合并文件内容
            for i, source_file in enumerate(files, 1):
                try:
                    with open(source_file.path, 'r', encoding='utf-8') as f:
                        html_content = f.read()

                    # 解析 HTML
                    file_soup = BeautifulSoup(html_content, 'html.parser')

                    # 添加分隔符和标题
                    section = soup.new_tag('section', **{'id': f'section-{i}'})
                    section_h2 = soup.new_tag('h2')
                    section_h2.string = source_file.path.stem
                    section.append(section_h2)

                    # 提取 body 内容
                    file_body = file_soup.find('body')
                    if file_body:
                        for element in file_body.contents:
                            section.append(element.copy())

                    hr = soup.new_tag('hr')
                    section.append(hr)

                    body.append(section)

                except Exception as e:
                    logger.error(f"读取 HTML 文件失败 {source_file.path}: {e}")
                    self.stats['errors'].append(str(source_file.path))

            # 保存文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(str(soup.prettify()))

            logger.info(f"已保存: {output_path}")
            return True

        except ImportError:
            logger.error("未安装 beautifulsoup4 库，请运行: pip3 install beautifulsoup4 lxml")
            return False
        except Exception as e:
            logger.error(f"合并 HTML 失败: {e}")
            return False

    def merge_epub(self, files: List[SourceFile], output_path: Path) -> bool:
        """
        合并 EPUB 文件（简单拼接）

        Args:
            files: 源文件信息列表
            output_path: 输出文件路径

        Returns:
            成功返回 True
        """
        try:
            import zipfile

            # 创建新的 EPUB
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as epub_out:
                # 添加基本的 EPUB 结构
                mimetype = b'application/epub+zip'
                epub_out.writestr('mimetype', mimetype)

                # 从第一个文件复制 META-INF 和 OEBPS
                if files:
                    try:
                        with zipfile.ZipFile(files[0].path, 'r') as epub_in:
                            for item in epub_in.namelist():
                                if not item.startswith('mimetype'):
                                    content = epub_in.read(item)
                                    epub_out.writestr(item, content)
                    except Exception as e:
                        logger.error(f"读取 EPUB 结构失败: {e}")

                # 合并内容文件（简化版）
                # 注意：这是基础实现，完整实现需要解析和重组 EPUB 的结构

            logger.info(f"已保存: {output_path} (基础合并)")
            return True

        except Exception as e:
            logger.error(f"合并 EPUB 失败: {e}")
            return False

    def merge_rtf(self, files: List[SourceFile], output_path: Path) -> bool:
        """
        合并 RTF 文件（文本提取）

        Args:
            files: 源文件信息列表
            output_path: 输出文件路径

        Returns:
            成功返回 True
        """
        try:
            # 简单实现：提取文本并重新打包为 RTF
            content_parts = []

            for source_file in files:
                try:
                    with open(source_file.path, 'r', encoding='utf-8', errors='ignore') as f:
                        # RTF 格式复杂，这里做简单处理
                        content = f.read()
                        # 移除 RTF 头部，保留文本部分
                        content = re.sub(r'\\[a-zA-Z]+\d*', '', content)
                        content_parts.append(f"## {source_file.name}\n\n{content}\n\n{self.separator}\n\n")
                except Exception as e:
                    logger.error(f"读取 RTF 文件失败 {source_file.path}: {e}")

            # 保存为文本（RTF 完整合并需要专门的库）
            txt_path = output_path.with_suffix('.txt')
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.writelines(content_parts)

            logger.info(f"已保存: {txt_path} (RTF 转文本)")
            return True

        except Exception as e:
            logger.error(f"合并 RTF 失败: {e}")
            return False

    def merge_odt(self, files: List[SourceFile], output_path: Path) -> bool:
        """
        合并 ODT 文件（使用 zipfile）

        Args:
            files: 源文件信息列表
            output_path: 输出文件路径

        Returns:
            成功返回 True
        """
        try:
            # ODT 是 ZIP 格式，这里做简单处理
            # 完整实现需要使用 odfpy 或类似库
            logger.warning("ODT 合并功能需要额外的库支持，建议转换为其他格式")

            # 简单实现：提取文本
            import zipfile

            content_parts = []

            for source_file in files:
                try:
                    with zipfile.ZipFile(source_file.path, 'r') as odt:
                        # 读取 content.xml
                        content_data = odt.read('content.xml')
                        # 简单的文本提取
                        content = re.sub(r'<[^>]+>', ' ', content_data.decode('utf-8'))
                        content = re.sub(r'\s+', ' ', content).strip()
                        content_parts.append(f"## {source_file.name}\n\n{content}\n\n{self.separator}\n\n")
                except Exception as e:
                    logger.error(f"读取 ODT 文件失败 {source_file.path}: {e}")

            # 保存为文本
            txt_path = output_path.with_suffix('.txt')
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.writelines(content_parts)

            logger.info(f"已保存: {txt_path} (ODT 转文本)")
            return True

        except Exception as e:
            logger.error(f"合并 ODT 失败: {e}")
            return False

    def merge_files(self, files: List[SourceFile], file_type: str, output_number: int) -> bool:
        """
        合并文件列表

        Args:
            files: 源文件信息列表
            file_type: 文件类型
            output_number: 输出文件编号

        Returns:
            成功返回 True
        """
        if not files:
            logger.warning("没有文件需要合并")
            return False

        # 生成输出文件名
        ext = self.FILE_TYPES[file_type]['ext']
        output_name = f"merged_{output_number}{ext}"
        output_path = self.output_dir / output_name

        logger.info(f"正在合并 {len(files)} 个文件到 {output_name}")

        # 计算总大小
        total_size = sum(f.size for f in files)

        # 创建输出文件信息
        output_info = OutputFile(
            name=output_name,
            source_files=files,
            total_size=total_size
        )
        self.report_outputs.append(output_info)

        # 根据文件类型调用相应的合并方法
        success = False
        if file_type == 'md':
            success = self.merge_markdown(files, output_path)
        elif file_type == 'txt':
            success = self.merge_text(files, output_path)
        elif file_type == 'pdf':
            success = self.merge_pdf(files, output_path)
        elif file_type == 'docx':
            success = self.merge_docx(files, output_path)
        elif file_type in ['html', 'htm']:
            success = self.merge_html(files, output_path)
        elif file_type == 'epub':
            success = self.merge_epub(files, output_path)
        elif file_type == 'rtf':
            success = self.merge_rtf(files, output_path)
        elif file_type == 'odt':
            success = self.merge_odt(files, output_path)
        else:
            logger.error(f"不支持的文件类型: {file_type}")
            return False

        if success:
            self.stats['outputs_created'] += 1

            # 标记文件为已处理
            if self.incremental:
                for source_file in files:
                    file_key = f"{source_file.path}:{source_file.hash}"
                    self.processed_files.add(file_key)

        return success

    def generate_merge_report(self) -> bool:
        """
        生成合并报告

        Returns:
            成功返回 True
        """
        if not self.generate_report:
            return False

        try:
            # 创建报告数据
            report = MergeReport(
                timestamp=datetime.now().isoformat(),
                input_directory=str(self.input_dir),
                output_directory=str(self.output_dir),
                file_type=self.file_type,
                scan_mode="递归" if self.recursive else "单层",
                total_files_found=self.stats['total_scanned'],
                total_files_processed=self.stats['total_processed'],
                outputs_created=self.stats['outputs_created'],
                balance_mode="按大小" if self.balance_by_size else "按数量",
                outputs=self.report_outputs,
                errors=self.stats['errors']
            )

            # 生成 JSON 报告
            if self.generate_report.endswith('.json') or not self.generate_report.endswith('.csv'):
                json_path = Path(self.generate_report)
                if not json_path.suffix:
                    json_path = json_path.with_suffix('.json')

                with open(json_path, 'w', encoding='utf-8') as f:
                    json.dump(asdict(report), f, ensure_ascii=False, indent=2, default=str)
                logger.info(f"JSON 报告已保存: {json_path}")

                # 同时生成 CSV 报告
                csv_path = json_path.with_suffix('.csv')
                self._generate_csv_report(report, csv_path)

            return True

        except Exception as e:
            logger.error(f"生成报告失败: {e}")
            return False

    def _generate_csv_report(self, report: MergeReport, csv_path: Path):
        """生成 CSV 格式的报告"""
        try:
            with open(csv_path, 'w', newline='', encoding='utf-8-sig') as f:
                writer = csv.writer(f)

                # 写入摘要信息
                writer.writerow(['合并报告摘要'])
                writer.writerow(['时间戳', report.timestamp])
                writer.writerow(['输入目录', report.input_directory])
                writer.writerow(['输出目录', report.output_directory])
                writer.writerow(['文件类型', report.file_type])
                writer.writerow(['扫描模式', report.scan_mode])
                writer.writerow(['发现文件数', report.total_files_found])
                writer.writerow(['处理文件数', report.total_files_processed])
                writer.writerow(['创建输出数', report.outputs_created])
                writer.writerow(['平衡模式', report.balance_mode])
                writer.writerow([])

                # 写入输出文件详情
                writer.writerow(['输出文件', '包含的源文件', '文件数量', '总大小(字节)'])

                for output in report.outputs:
                    source_names = ', '.join([f.name for f in output.source_files])
                    writer.writerow([
                        output.name,
                        source_names,
                        len(output.source_files),
                        output.total_size
                    ])

                # 写入错误信息
                if report.errors:
                    writer.writerow([])
                    writer.writerow(['错误信息'])
                    for error in report.errors:
                        writer.writerow([error])

            logger.info(f"CSV 报告已保存: {csv_path}")

        except Exception as e:
            logger.error(f"生成 CSV 报告失败: {e}")

    def process(self, dry_run: bool = False) -> bool:
        """
        执行合并流程

        Args:
            dry_run: 预览模式，不实际合并

        Returns:
            成功返回 True
        """
        try:
            # 扫描文件
            files_by_type = self.scan_files()

            if not files_by_type:
                logger.error("未找到任何可处理的文件")
                return False

            # 加载排序文件
            order = self.load_order_file()

            # 处理每种文件类型
            for file_type, files in files_by_type.items():
                logger.info(f"\n处理 {file_type.upper()} 文件...")

                # 应用文件数量限制
                if self.total_files and self.total_files < len(files):
                    files = files[:self.total_files]
                    logger.info(f"限制处理前 {self.total_files} 个文件")

                # 排序文件
                sorted_files = self.sort_files(files, order)

                if not sorted_files:
                    logger.warning(f"没有 {file_type} 文件需要处理")
                    continue

                # 计算分配
                allocation = self.calculate_allocation(sorted_files)

                logger.info(f"\n合并计划:")
                logger.info(f"  源文件数: {len(sorted_files)}")
                logger.info(f"  输出文件数: {len(allocation)}")
                if self.balance_by_size:
                    total_size = sum(f.size for f in sorted_files)
                    logger.info(f"  总大小: {total_size} 字节")

                if dry_run:
                    logger.info(f"\n预览模式（不会实际合并）:")
                    for i, (start, end) in enumerate(allocation, 1):
                        file_list = sorted_files[start:end]
                        logger.info(f"\n  [输出 {i}] merged_{i}{self.FILE_TYPES[file_type]['ext']}:")
                        for f in file_list:
                            size_mb = f.size / 1024 / 1024
                            logger.info(f"    - {f.name} ({size_mb:.2f} MB)")
                    continue

                # 执行合并
                for i, (start, end) in enumerate(allocation, 1):
                    file_list = sorted_files[start:end]
                    logger.info(f"\n[{i}/{len(allocation)}] 处理输出文件...")

                    for f in file_list:
                        size_kb = f.size / 1024
                        logger.info(f"  - {f.name} ({size_kb:.1f} KB)")

                    self.merge_files(file_list, file_type, i)
                    self.stats['total_processed'] += len(file_list)

            # 保存状态
            if self.incremental:
                self._save_state()

            # 生成报告
            if not dry_run:
                self.generate_merge_report()

            # 打印统计信息
            self.print_stats()

            return True

        except Exception as e:
            logger.error(f"处理失败: {e}")
            return False

    def print_stats(self):
        """打印统计信息"""
        print("\n" + "=" * 50)
        print("合并统计")
        print("=" * 50)
        print(f"扫描文件: {self.stats['total_scanned']}")
        print(f"处理文件: {self.stats['total_processed']}")
        print(f"创建输出: {self.stats['outputs_created']}")
        if self.stats['errors']:
            print(f"错误数量: {len(self.stats['errors'])}")
            print("\n错误文件:")
            for error in self.stats['errors'][:10]:
                print(f"  - {error}")
            if len(self.stats['errors']) > 10:
                print(f"  ... 还有 {len(self.stats['errors']) - 10} 个")
        print(f"\n输出目录: {self.output_dir}")
        print("=" * 50)


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description="文件批量合并工具（增强版）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 基本合并
  %(prog)s --input-dir ./docs --output-dir ./merged \\
           --file-type md --output-count 10

  # 递归扫描子目录
  %(prog)s --input-dir ./docs --output-dir ./merged \\
           --file-type md --recursive --output-count 5

  # 使用正则表达式过滤
  %(prog)s --input-dir ./docs --output-dir ./merged \\
           --file-type md --filter "chapter_.*" --output-count 10

  # 按文件大小平衡分配
  %(prog)s --input-dir ./docs --output-dir ./merged \\
           --file-type md --balance-by-size --output-count 5

  # 生成合并报告
  %(prog)s --input-dir ./docs --output-dir ./merged \\
           --file-type md --output-count 10 --generate-report report.json

  # 增量合并模式
  %(prog)s --input-dir ./docs --output-dir ./merged \\
           --file-type md --incremental --state-file state.json \\
           --output-count 10

  # 组合使用
  %(prog)s --input-dir ./docs --output-dir ./merged \\
           --file-type all --recursive --balance-by-size \\
           --generate-report report.json --output-count 10
        """
    )

    parser.add_argument('--input-dir', required=True, help='输入目录路径')
    parser.add_argument('--output-dir', required=True, help='输出目录路径')
    parser.add_argument('--file-type', required=True,
                       choices=['md', 'pdf', 'docx', 'txt', 'epub', 'html', 'rtf', 'odt', 'all'],
                       help='文件类型')
    parser.add_argument('--output-count', type=int, required=True,
                       help='要生成的输出文件数量')
    parser.add_argument('--order-file', help='自定义排序文件路径')
    parser.add_argument('--total-files', type=int, help='指定要处理的文件总数')
    parser.add_argument('--separator', default='---', help='文件分隔符（默认：---）')
    parser.add_argument('--no-index', action='store_true', help='不生成目录索引')
    parser.add_argument('--dry-run', action='store_true', help='预览合并计划，不实际执行')

    # 扩展功能参数
    parser.add_argument('--recursive', action='store_true',
                       help='递归扫描子目录')
    parser.add_argument('--filter', help='包含文件的正则表达式')
    parser.add_argument('--exclude', help='排除文件的正则表达式')
    parser.add_argument('--balance-by-size', action='store_true',
                       help='按文件大小平衡分配（而非按数量）')
    parser.add_argument('--generate-report', help='生成合并报告（JSON/CSV）')
    parser.add_argument('--incremental', action='store_true',
                       help='增量合并模式（只处理新文件）')
    parser.add_argument('--state-file', help='状态文件路径（用于增量合并）')

    args = parser.parse_args()

    # 创建合并器
    merger = FileMerger(
        input_dir=args.input_dir,
        output_dir=args.output_dir,
        file_type=args.file_type,
        output_count=args.output_count,
        order_file=args.order_file,
        total_files=args.total_files,
        separator=args.separator,
        generate_index=not args.no_index,
        recursive=args.recursive,
        filter_pattern=args.filter,
        exclude_pattern=args.exclude,
        balance_by_size=args.balance_by_size,
        generate_report=args.generate_report,
        incremental=args.incremental,
        state_file=args.state_file
    )

    # 执行合并
    success = merger.process(dry_run=args.dry_run)

    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
